import os
import logging
from docx import Document
from weasyprint import HTML, CSS
from PyPDF2 import PdfWriter, PdfReader
import io

from utils.config import OUTPUTS_DIR

logger = logging.getLogger(__name__)

def create_docx_report(user_id: str, title: str, content: str, filename: str) -> str:
    """
    Creates a .docx file with the given title and content.
    """
    if not filename.endswith(".docx"):
        filename += ".docx"

    output_path = os.path.join(OUTPUTS_DIR, filename)
    logger.info(f"Creating DOCX report at: {output_path}")

    try:
        document = Document()
        document.add_heading(title, level=1)

        # Add content paragraphs
        for paragraph in content.split('\n'):
            if paragraph.strip(): # Avoid adding empty paragraphs
                document.add_paragraph(paragraph)

        document.save(output_path)
        logger.info(f"Successfully created DOCX report: {filename}")
        return f"Successfully created DOCX report named '{filename}'."
    except Exception as e:
        logger.error(f"Error creating DOCX report '{filename}': {e}", exc_info=True)
        return f"Error: Failed to create DOCX report. {e}"

def create_pdf_report(user_id: str, title: str, content: str, filename: str) -> str:
    """
    Creates a .pdf file from the given title and content using HTML as an intermediate.
    """
    if not filename.endswith(".pdf"):
        filename += ".pdf"

    output_path = os.path.join(OUTPUTS_DIR, filename)
    logger.info(f"Creating PDF report at: {output_path}")

    try:
        # Convert markdown-like newlines to HTML paragraphs
        html_content = "".join(f"<p>{p.strip()}</p>" for p in content.split('\n') if p.strip())
        html_string = f"""
        <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: sans-serif; }}
                    h1 {{ color: #333; }}
                    p {{ line-height: 1.6; }}
                </style>
            </head>
            <body>
                <h1>{title}</h1>
                {html_content}
            </body>
        </html>
        """

        # Use WeasyPrint to render HTML to PDF
        HTML(string=html_string).write_pdf(output_path)

        logger.info(f"Successfully created PDF report: {filename}")
        return f"Successfully created PDF report named '{filename}'."
    except Exception as e:
        logger.error(f"Error creating PDF report '{filename}': {e}", exc_info=True)
        return f"Error: Failed to create PDF report. {e}"

